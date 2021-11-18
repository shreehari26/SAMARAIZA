import os
import wave
import time
import pickle
import pyaudio
import warnings
import numpy as np
from sklearn import preprocessing
from scipy.io.wavfile import read
import python_speech_features as mfcc
from sklearn.mixture import GaussianMixture
import speech_recognition as sr
import docx
import yake
import spacy

from chatbot2 import *

# !python -m spacy download en_core_web_sm

from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation

import sumy
import nltk
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer

from sumy.nlp.tokenizers import Tokenizer
from sumy.parsers.plaintext import PlaintextParser

warnings.filterwarnings("ignore")

list_req_pp = ["meet2","rohit","sameer"]

def calculate_delta(array):
    rows, cols = array.shape
    print(rows)
    print(cols)
    deltas = np.zeros((rows, 20))
    N = 2
    for i in range(rows):
        index = []
        j = 1
        while j <= N:
            if i - j < 0:
                first = 0
            else:
                first = i - j
            if i + j > rows - 1:
                second = rows - 1
            else:
                second = i + j
            index.append((second, first))
            j += 1
        deltas[i] = (array[index[0][0]] - array[index[0][1]] + (2 * (array[index[1][0]] - array[index[1][1]]))) / 10
    return deltas


def extract_features(audio, rate):
    mfcc_feature = mfcc.mfcc(audio, rate, 0.025, 0.01, 20, nfft=1200, appendEnergy=True)
    mfcc_feature = preprocessing.scale(mfcc_feature)
    print(mfcc_feature)
    delta = calculate_delta(mfcc_feature)
    combined = np.hstack((mfcc_feature, delta))
    return combined


def record_audio_train():

    Name = (input("Please Enter Your Name:"))
    # record_audio_train.name_train = Name
    list_req_pp.append(Name)
    for count in range(5):
        FORMAT = pyaudio.paInt16
        CHANNELS = 1
        RATE = 44100
        CHUNK = 512
        RECORD_SECONDS = 10
        device_index = 2
        audio = pyaudio.PyAudio()
        print("----------------------record device list---------------------")
        info = audio.get_host_api_info_by_index(0)
        numdevices = info.get('deviceCount')
        for i in range(0, numdevices):
            if (audio.get_device_info_by_host_api_device_index(0, i).get('maxInputChannels')) > 0:
                print("Input Device id ", i, " - ", audio.get_device_info_by_host_api_device_index(0, i).get('name'))
        print("-------------------------------------------------------------")
        index = int(input())
        print("recording via index " + str(index))
        stream = audio.open(format=FORMAT, channels=CHANNELS,
                            rate=RATE, input=True, input_device_index=index,
                            frames_per_buffer=CHUNK)
        print("recording started")
        Recordframes = []
        for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
            data = stream.read(CHUNK)
            Recordframes.append(data)
        print("recording stopped")
        stream.stop_stream()
        stream.close()
        audio.terminate()
        OUTPUT_FILENAME = Name + "-sample" + str(count) + ".wav"
        WAVE_OUTPUT_FILENAME = os.path.join("training_set", OUTPUT_FILENAME)
        trainedfilelist = open("training_set_addition.txt", 'a')
        trainedfilelist.write(OUTPUT_FILENAME + "\n")
        waveFile = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
        waveFile.setnchannels(CHANNELS)
        waveFile.setsampwidth(audio.get_sample_size(FORMAT))
        waveFile.setframerate(RATE)
        waveFile.writeframes(b''.join(Recordframes))
        waveFile.close()


def record_audio_test():
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 44100
    CHUNK = 512
    RECORD_SECONDS = 10
    device_index = 2
    audio = pyaudio.PyAudio()
    print("----------------------record device list---------------------")
    info = audio.get_host_api_info_by_index(0)
    numdevices = info.get('deviceCount')
    for i in range(0, numdevices):
        if (audio.get_device_info_by_host_api_device_index(0, i).get('maxInputChannels')) > 0:
            print("Input Device id ", i, " - ", audio.get_device_info_by_host_api_device_index(0, i).get('name'))
    print("-------------------------------------------------------------")
    index = int(input())
    print("recording via index " + str(index))
    stream = audio.open(format=FORMAT, channels=CHANNELS,
                        rate=RATE, input=True, input_device_index=index,
                        frames_per_buffer=CHUNK)
    print("recording started")
    Recordframes = []
    for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
        data = stream.read(CHUNK)
        Recordframes.append(data)
    print("recording stopped")
    stream.stop_stream()
    stream.close()
    audio.terminate()

    OUTPUT_FILENAME = "sample.wav"
    WAVE_OUTPUT_FILENAME = os.path.join("testing_set", OUTPUT_FILENAME)
    trainedfilelist = open("testing_set_addition.txt", 'a')
    trainedfilelist.write(OUTPUT_FILENAME + "\n")
    waveFile = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
    waveFile.setnchannels(CHANNELS)
    waveFile.setsampwidth(audio.get_sample_size(FORMAT))
    waveFile.setframerate(RATE)
    waveFile.writeframes(b''.join(Recordframes))
    waveFile.close()


def train_model():
    source = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\training_set\\"
    dest = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\trained_models\\"
    train_file = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\training_set_addition.txt"
    file_paths = open(train_file, 'r')
    count = 1
    features = np.asarray(())
    for path in file_paths:
        path = path.strip()
        print(path)

        sr, audio = read(source + path)
        print(sr)
        vector = extract_features(audio, sr)

        if features.size == 0:
            features = vector
        else:
            features = np.vstack((features, vector))

        if count == 5:
            gmm = GaussianMixture(n_components=6, max_iter=200, covariance_type='diag', n_init=3)
            gmm.fit(features)

            # dumping the trained gaussian model
            picklefile = path.split("-")[0] + ".gmm"
            pickle.dump(gmm, open(dest + picklefile, 'wb'))
            print('+ modeling completed for speaker:', picklefile, " with data point = ", features.shape)
            features = np.asarray(())
            count = 0
        count = count + 1


def test_model():
    source = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\testing_set\\"
    modelpath = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\trained_models\\"
    test_file = "C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\testing_set_addition.txt"
    file_paths = open(test_file, 'r')

    gmm_files = [os.path.join(modelpath, fname) for fname in
                 os.listdir(modelpath) if fname.endswith('.gmm')]

    # Load the Gaussian gender Models
    models = [pickle.load(open(fname, 'rb')) for fname in gmm_files]
    speakers = [fname.split("\\")[-1].split(".gmm")[0] for fname
                in gmm_files]

    # Read the test directory and get the list of test audio files
    for path in file_paths:

        path = path.strip()
        print(path)
        sr, audio = read(source + path)
        vector = extract_features(audio, sr)

        log_likelihood = np.zeros(len(models))

        for i in range(len(models)):
            gmm = models[i]  # checking with each model one by one
            scores = np.array(gmm.score(vector))
            log_likelihood[i] = scores.sum()
        print(log_likelihood)
        winner = np.argmax(log_likelihood)

        print("\tdetected as - ", speakers[winner])
        test_model.speaker_var = speakers[winner]




        time.sleep(1.0)


# choice=int(input("\n1.Record audio for training \n 2.Train Model \n 3.Record audio for testing \n 4.Test Model\n"))

while True:
    choice = int(
        input("\n 1.Record audio for training \n 2.Train Model \n 3.Record audio for testing \n 4.Test Model\n 5.chatbot"))
    if (choice == 1):
        record_audio_train()
    elif (choice == 2):
        train_model()
    elif (choice == 3):
        record_audio_test()
    elif (choice == 4):
        test_model()
    elif (choice==5):
        chat()


        if test_model.speaker_var in list_req_pp:
            print('correct')


            def audio_to_text():
                Name = (input("Please Enter Your Name:"))
                audio_to_text.audio_file_name = Name
                for count in range(1):
                    FORMAT = pyaudio.paInt16
                    CHANNELS = 1
                    RATE = 44100
                    CHUNK = 512
                    RECORD_SECONDS = 60##user input
                    device_index = 2
                    audio = pyaudio.PyAudio()
                    print("----------------------record device list---------------------")
                    info = audio.get_host_api_info_by_index(0)
                    numdevices = info.get('deviceCount')
                    for i in range(0, numdevices):
                        if (audio.get_device_info_by_host_api_device_index(0, i).get('maxInputChannels')) > 0:
                            print("Input Device id ", i, " - ",
                                  audio.get_device_info_by_host_api_device_index(0, i).get('name'))
                    print("-------------------------------------------------------------")
                    index = int(input())
                    print("recording via index " + str(index))
                    stream = audio.open(format=FORMAT, channels=CHANNELS,
                                        rate=RATE, input=True, input_device_index=index,
                                        frames_per_buffer=CHUNK)
                    print("recording started")
                    Recordframes = []
                    for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
                        data = stream.read(CHUNK)
                        Recordframes.append(data)
                    print("recording stopped")

                    stream.stop_stream()
                    stream.close()
                    audio.terminate()
                    OUTPUT_FILENAME = Name + "check1" + str(count) + ".wav"
                    WAVE_OUTPUT_FILENAME = os.path.join("text_audio", OUTPUT_FILENAME)

                    audio_to_text.file_name = OUTPUT_FILENAME

                    waveFile = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
                    waveFile.setnchannels(CHANNELS)
                    waveFile.setsampwidth(audio.get_sample_size(FORMAT))
                    waveFile.setframerate(RATE)
                    waveFile.writeframes(b''.join(Recordframes))
                    waveFile.close()


            audio_to_text()

            sound = f"C:\\Users\\Rohit\\Dropbox\\My PC (DESKTOP-BV3E4EV)\\Desktop\\spec_rec3\\text_audio\\{audio_to_text.file_name}"

            r = sr.Recognizer()

            with sr.AudioFile(sound) as source:
                r.adjust_for_ambient_noise(source)

                print("Converting Audio File To Text...")

                audio = r.listen(source)
#sumariser
                try:
                    text = str(r.recognize_google(audio))
                    choice = int(input("1)trim according frequecy\n2)summarise to required number of lines\n3)summarise to core message"))

                    if choice==1:##FREQUENCY
                        stopwords = list(STOP_WORDS)
                        nlp = spacy.load('en_core_web_sm')
                        doc = nlp(text)
                        tokens = [token.text for token in doc]

                        a = text.split(".")
                        text = " ".join(a)

                        punctuation = punctuation + '\n'

                        word_frequencies = {}
                        for word in doc:
                            if word.text.lower() not in stopwords:
                                if word.text.lower() not in punctuation:
                                    if word.text not in word_frequencies.keys():
                                        word_frequencies[word.text] = 1
                                    else:
                                        word_frequencies[word.text] += 1

                        max_frequency = max(word_frequencies.values())
                        for word in word_frequencies.keys():
                            word_frequencies[word] = word_frequencies[word] / max_frequency

                        sentence_tokens = [sent for sent in doc.sents]
                        print(sentence_tokens)

                        sentence_scores = {}
                        for sent in sentence_tokens:
                            for word in sent:
                                if word.text.lower() in word_frequencies.keys():
                                    if sent not in sentence_scores.keys():
                                        sentence_scores[sent] = word_frequencies[word.text.lower()]
                                    else:
                                        sentence_scores[sent] += word_frequencies[word.text.lower()]

                        from heapq import nlargest

                        select_length = int(len(sentence_tokens) * 0.3)

                        summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)

                        final_summary = [word.text for word in summary]
                        summary = ' '.join(final_summary)

                        print(summary)

                        kw_extractor = yake.KeywordExtractor()
                        language = "en"
                        max_ngram_size = 3
                        deduplication_threshold = 0.9
                        numOfKeywords = 20
                        custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                                                                    dedupLim=deduplication_threshold, top=numOfKeywords,
                                                                    features=None)
                        keywords = custom_kw_extractor.extract_keywords(text)
                        keywords_list = []
                        for kw in keywords:
                            keywords_list.append(kw[0])

                        doc = docx.Document()

                        # add a heading of level 0 (largest heading)
                        doc.add_heading('Heading for the document', 0)

                        # add a paragraph and store
                        # the object in a variable
                        doc_para = doc.add_paragraph(summary)

                        # add a run i.e, style like
                        # bold, italic, underline, etc.

                        for bolding in keywords_list:
                            doc_para.add_run(bolding).bold = True

                        # doc_para.add_run(', and ')
                        # doc_para.add_run('these words are italic').italic = True

                        # add a page break to start a new page
                        # doc.add_page_break()

                        # add a heading of level 2
                        # doc.add_heading('Heading level 2', 2)

                        # pictures can also be added to our word document
                        # width is optional
                        # doc.add_picture('path_to_picture')

                        # now save the document to a location
                        doc.save(audio_to_text.audio_file_name)



                    elif choice==2:#REQUIRED NUMBER OF SENTENCES
                        check1 = text.split(" ")
                        # print(check1)
                        check2 = []
                        c = 0
                        for i in check1:
                            c += 1
                            if c == 15:
                                check2.append(".")
                                c = 0
                            else:
                                check2.append(i)

                        # print(check2)
                        text2 = " ".join(check2)
                        # print(text2)
                        parser = PlaintextParser.from_string(text2, Tokenizer('english'))

                        lsa_summarizer = LsaSummarizer()
                        lsa_summary = lsa_summarizer(parser.document, 3)

                        a = ''

                        for sentence in lsa_summary:
                            # print(sentence)
                            a = a + str(sentence)

                        # print(len(text.split(' ')))
                        # print(len(a.split(' ')))
                        sentences_text = a.split('.')
                        for sent in sentences_text:
                            print(sent)
                        # kw_extractor = yake.KeywordExtractor()
                        # language = "en"
                        # max_ngram_size = 3
                        # deduplication_threshold = 0.9
                        # numOfKeywords = 20
                        # custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                        #                                             dedupLim=deduplication_threshold, top=numOfKeywords,
                        #                                             features=None)
                        # keywords = custom_kw_extractor.extract_keywords(text)
                        # keywords_list = []
                        # for kw in keywords:
                        #     keywords_list.append(kw[0])
                        #
                        # doc = docx.Document()
                        #
                        # # add a heading of level 0 (largest heading)
                        # doc.add_heading('Heading for the document', 0)
                        #
                        # # add a paragraph and store
                        # # the object in a variable
                        # doc_para = doc.add_paragraph(summary)
                        #
                        # # add a run i.e, style like
                        # # bold, italic, underline, etc.
                        #
                        # for bolding in keywords_list:
                        #     doc_para.add_run(bolding).bold = True
                        #
                        # # doc_para.add_run(', and ')
                        # # doc_para.add_run('these words are italic').italic = True
                        #
                        # # add a page break to start a new page
                        # # doc.add_page_break()
                        #
                        # # add a heading of level 2
                        # # doc.add_heading('Heading level 2', 2)
                        #
                        # # pictures can also be added to our word document
                        # # width is optional
                        # # doc.add_picture('path_to_picture')
                        #
                        # # now save the document to a location
                        # doc.save(audio_to_text.audio_file_name)
                    # elif choice==3:
#                        core message







                except Exception as e:
                    print(e)

    if (choice > 4):
        exit()